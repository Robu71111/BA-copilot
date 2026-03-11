"""
Document Parser — v2
====================
Parses .txt / .docx / .pdf files and returns plain text.

Key fixes vs v1:
  - Replaces PyPDF2 (abandoned) with pypdf (actively maintained)
  - TXT: tries utf-8, falls back to latin-1, then chardet
  - DOCX: also extracts text from tables (v1 missed them)
  - PDF: uses pypdf PdfReader; falls back page-by-page on extraction errors
  - All methods return empty string rather than raising on empty pages

Install / requirements.txt additions:
    pypdf>=4.0.0        (replaces PyPDF2)
    python-docx>=1.1.0  (unchanged)
"""

import io
import logging

logger = logging.getLogger(__name__)


# ── Lazy imports so startup never crashes on missing libs ─────────────────────
def _get_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Add 'python-docx>=1.1.0' to requirements.txt"
        )


def _get_pdf_reader():
    # Try pypdf first (modern), fall back to PyPDF2 (old) so existing deploys keep working
    try:
        from pypdf import PdfReader
        return PdfReader
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader  # noqa: N813
        logger.warning("Using deprecated PyPDF2 — please upgrade to pypdf>=4.0.0")
        return PdfReader
    except ImportError:
        raise ImportError(
            "No PDF library found. Add 'pypdf>=4.0.0' to requirements.txt"
        )


# ── Parser ────────────────────────────────────────────────────────────────────
class DocumentParser:
    """Parse various document formats and extract plain text."""

    @staticmethod
    def parse_txt(file_content: bytes | str) -> str:
        """
        Parse a plain-text file.
        Tries UTF-8 → latin-1 → chardet as fallback encodings.
        """
        if isinstance(file_content, str):
            return file_content.strip()

        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return file_content.decode(enc).strip()
            except (UnicodeDecodeError, LookupError):
                continue

        # Last resort: chardet
        try:
            import chardet
            detected = chardet.detect(file_content)
            enc = detected.get("encoding") or "utf-8"
            return file_content.decode(enc, errors="replace").strip()
        except Exception:
            pass

        return file_content.decode("utf-8", errors="replace").strip()

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Parse a .docx file.
        Extracts paragraphs AND table cells (v1 missed tables).
        """
        Document = _get_docx()
        doc = Document(io.BytesIO(file_content))

        parts: list[str] = []

        # Body paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Table cells — iterate all tables in the document
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        if not parts:
            raise ValueError("No readable text found in the DOCX file.")

        return "\n\n".join(parts).strip()

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Parse a PDF file using pypdf (or PyPDF2 as fallback).
        Extracts text page by page; skips pages that fail individually.
        """
        PdfReader = _get_pdf_reader()
        reader = PdfReader(io.BytesIO(file_content))

        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    pages.append(text)
            except Exception as exc:
                logger.warning(f"[doc_parser] Skipping PDF page {i}: {exc}")
                continue

        if not pages:
            raise ValueError(
                "No readable text found in the PDF. "
                "If it is a scanned PDF, OCR support is required."
            )

        return "\n\n--- Page Break ---\n\n".join(pages).strip()

    @staticmethod
    def parse_document(file_name: str, file_content: bytes) -> str:
        """
        Auto-detect format from file extension and parse accordingly.
        Returns extracted plain text.
        """
        ext = file_name.lower().rsplit(".", 1)[-1]

        if ext == "txt":
            return DocumentParser.parse_txt(file_content)
        elif ext == "docx":
            return DocumentParser.parse_docx(file_content)
        elif ext == "pdf":
            return DocumentParser.parse_pdf(file_content)
        else:
            raise ValueError(
                f"Unsupported file format: .{ext}. "
                "Supported formats: .txt, .docx, .pdf"
            )

    @staticmethod
    def validate_text(text: str, min_length: int = 50) -> tuple[bool, str]:
        """
        Validate that extracted text meets minimum quality bar.
        Returns (is_valid, message).
        """
        if not text or not text.strip():
            return False, "No text content found in the document."
        clean = text.strip()
        if len(clean) < min_length:
            return False, (
                f"Document text is too short ({len(clean)} chars). "
                f"Minimum {min_length} characters required."
            )
        return True, "Text validation successful."


# Singleton — imported by routes as `from backend.services.document_parser import parser`
parser = DocumentParser()